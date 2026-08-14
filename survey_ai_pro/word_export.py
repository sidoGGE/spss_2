"""📝 تصدير Word احترافي مع دعم الصور والشروح"""
from pathlib import Path
from typing import Dict, List, Optional
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rich.console import Console

console = Console()

class WordExporter:
    def __init__(self, path: str = "output/report.docx"):
        self.path = Path(path)
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    def _shade(self, cell, color: str):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shd)

    def title(self, text: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        p.space_after = Pt(16)

    def heading(self, text: str, level: int = 1):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
        r.font.color.rgb = RGBColor(0xA2, 0x3B, 0x72) if level <= 2 else RGBColor(0x46, 0x86, 0xAB)
        p.space_before = Pt(14 if level == 1 else 10)
        p.space_after = Pt(8)

    def paragraph(self, text: str, bold: bool = False, align: str = "right"):
        p = self.doc.add_paragraph()
        if align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        r = p.add_run(text)
        r.font.size = Pt(12)
        r.bold = bold
        p.space_after = Pt(8)

    def caption(self, text: str):
        """تسمية توضيحية للصور والجداول"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.space_after = Pt(10)

    def table(self, df: pd.DataFrame, title: str = None):
        if title:
            self.heading(title, level=3)

        t = self.doc.add_table(rows=1, cols=len(df.columns))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = 'Table Grid'

        hdr = t.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
            self._shade(hdr[i], "2E86AB")
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)

        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = f"{val:.3f}" if isinstance(val, float) else str(val)
                for para in cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
        self.doc.add_paragraph()

    def image(self, path: str, width: float = 15):
        if Path(path).exists():
            self.doc.add_picture(str(path), width=Cm(width))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph()

    def page_break(self):
        self.doc.add_page_break()

    def ai_text(self, text: str):
        self.heading("فصل نتائج الدراسة (مولد بالذكاء الاصطناعي)", level=1)
        self.paragraph(
            "يُقدم القسم التالي تحليلاً أكاديمياً شاملاً لمناقشة النتائج الإحصائية، "
            "مولداً بالذكاء الاصطناعي بناءً على البيانات الفعلية للدراسة."
        )

        for para in text.split('\n\n'):
            para = para.strip()
            if not para:
                continue
            if para.startswith('▪') or 'القسم' in para or 'المطلوب' in para or para.startswith('---'):
                self.heading(para.replace('▪', '').strip(), level=2)
            elif para.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                self.paragraph(para, bold=False)
            elif len(para) < 60:
                self.heading(para, level=3)
            else:
                self.paragraph(para)

    def save(self) -> str:
        self.doc.save(str(self.path))
        console.print(f"[green]✅ Word: {self.path}[/green]")
        return str(self.path)
